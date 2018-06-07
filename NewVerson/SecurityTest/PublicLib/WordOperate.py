#-*- coding: utf-8 -*-

import os
import sys
import docx
import traceback
from win32com.client import Dispatch,constants

reload(sys)
sys.setdefaultencoding('utf-8')

import GlobalValue as g_Global
import Log
g_Log = None
import LocalOperate
g_Local = None

class Word:
    def __init__(self,docName=None):
        global g_Log,g_Local
        g_Log = Log.Log()
        g_Local = LocalOperate.Local()
        self.docName = docName

    def new(self):#新建一个word
        try:
            g_Log.writeLog("该功能暂未实现")
            return 1
            # 检查excel所在路径是否存在，如不存在则新建完整目录
            excelPath = os.path.split(self.excelName)[0]
            pathExist = os.path.exists(excelPath)
            if not pathExist:
                os.makedirs(excelPath)
        except:
            g_Log.writeLog("traceback")
            return 0

    def write(self,result,redLine=0):# 向word写入
        try:
            g_Log.writeLog("该功能暂未实现")
            return 1
            # 读取word
        except:
            g_Log.writeLog("traceback")
            return 0

    def readlines(self): #读取word内容
        output = []
        try:
            docName = self.docName
            if docName.endswith('.doc'):
                g_Log.writeLog(u"暂不支持.doc格式，请另存为.docx格式后重新执行")
                return 0
            file = self.Document()
            for i in range(len(file.paragraphs)):
                output.append(file.paragraphs[i].text)
            return output
        except:
            g_Log.writeLog("traceback")
            return 0


    def Document(self):
        try:
            file = docx.Document(self.docName)
        except:
            errmsg1 = ''.join(traceback.format_exception(*sys.exc_info()))
            try:
                file = docx.Document(r'{wordPath}'.format(wordPath=self.docName))
            except:
                errmsg2 = ''.join(traceback.format_exception(*sys.exc_info()))
                g_Log.writeLog(errmsg1)
                g_Log.writeLog(errmsg2)
        return file







