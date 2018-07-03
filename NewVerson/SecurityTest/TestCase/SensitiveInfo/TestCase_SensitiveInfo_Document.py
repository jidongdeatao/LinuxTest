#-*- coding: utf-8 -*-

import sys
import os
import traceback
import datetime
import re
import time
import shutil
import docx

from sys import path
# 系统默认Unicode解码，需要换成utf-8形式
reload(sys)
sys.setdefaultencoding('utf-8')

'''
""" 相关安全要求说明 """
《01 产品网络安全红线落地解读及指导V2_1.xls》：
        8.1.1 系统自身操作维护类口令满足“口令安全要求”
《中央软件院网络安全测试基线V2.0.xlsx》 中关于弱口令词典等要求
'''
'''
""" 脚本功能 """
扫描资料中的弱口令等问题点（目前只支持扫描word和excel），扫描策略见本脚本所在目录下的scanSensitivePolicy.xlsx
注意：不支持.doc文件，请打开后另存为.docx格式
'''
'''
""" 脚本配置执行说明 """
配置1：本脚本所在目录下的scanSensitivePolicy.xlsx，“policy”页。
配置2：本脚本所在目录下的scanSensitivePolicy.xlsx，“config”页。每一行配置都有说明，请仔细阅读。
'''

'''
""" 可以在此处下方添加自己的代码（函数） """
'''
try:
    g_Log = None
    g_Global = None
    g_caseName = None

    curFile = os.path.abspath(sys._getframe(0).f_code.co_filename)
    g_caseName = curFile.replace("\\","/")
    g_curDir = os.path.split(g_caseName)[0]
    path.append( g_caseName.split("TestCase")[0]+"PublicLib" )

    import GlobalValue as g_Global
    g_Global.init()
    g_Global.setValue("startTime",str(datetime.datetime.now()))

    import Log
    import ExcelOperate
    import WordOperate
    import ContainerOperate
    import LinuxOperate
    import LocalOperate
    g_Log = Log.Log()
    g_Local = LocalOperate.Local()

    ##### 获取环境配置信息
    excelName = g_caseName.split("TestCase")[0]+"Config/config.xlsx"
    excel0 = ExcelOperate.Excel(excelName=excelName,sheetName="vmInfo")
    g_vmInfo = excel0.read()
    del g_vmInfo[0]
    scanPolicyExcel = g_curDir + "/scanSensitivePolicy.xlsx"
    excel1 = ExcelOperate.Excel(excelName=scanPolicyExcel,sheetName="policy")
    g_scanPolicy = excel1.read()
    del g_scanPolicy[0]
    excel2 = ExcelOperate.Excel(excelName=scanPolicyExcel,sheetName="config")
    g_config = excel2.read()
    del g_config[0]
    g_ResultFile = g_caseName.replace("TestCase","Report").replace(".py","")+"-"+g_Global.getValue("startTime").replace(":","").replace(" ","").replace("-","")+".xls"
except:
    errmsg = ''.join(traceback.format_exception(*sys.exc_info()))
    print errmsg
    exit(0)

def getScanPolicy(): #获取所有支持系统扫描的策略信息
    global g_scanPolicy
    policy = []
    for p in g_scanPolicy:
        supportSys = p[6].lower()
        if p[0].lower() == "true" and "document" in supportSys.lower():
            policy.append(p)
    g_scanPolicy = policy
    return 1

def scanWord(matchKey,wordFile):
    result = []
    try:
        word = WordOperate.Word(docName=wordFile)
        wordRead = word.readlines()
        lineNum = len(wordRead)
        for row in range(1,lineNum):
            lineNum = str(row+1)
            lineInfo = wordRead[row]

            x1 = re.findall(matchKey,lineInfo,re.I)
            if x1 == []:
                continue
            result.append([wordFile,lineNum,lineInfo])
    except:
        g_Log.writeLog("traceback")
    return result

def scanExcel(matchKey,excelFile):
    result = []
    try:
        excel0 = ExcelOperate.Excel(excelName=excelFile)
        sheetNumber = excel0.sheetNumber()
        for i in range(0,sheetNumber):
            excel = ExcelOperate.Excel(excelName=excelFile,sheetID=i)
            excelInfo = excel.read()
            rowNum = len(excelInfo)
            for row in range(0,rowNum):
                colNum = len(excelInfo[row])
                for col in range(0,colNum):
                    position = excel.position(col,row)
                    position = " ["+"sheetID="+str(i)+";"+position+"] "
                    lineInfo = excelInfo[row][col]
                    if type(lineInfo) != type("str"):
                        lineInfo=str(lineInfo)
                    x1 = re.findall(matchKey,lineInfo,re.I)
                    if x1 == []:
                        continue
                    result.append([excelFile,position,lineInfo])
    except:
        g_Log.writeLog("traceback")
    return result

def scanDocument(reportPath,file):
    for policy in g_scanPolicy:
        matchKey = policy[2]
        ifRed = policy[7]
        result = []
        if ifRed == "YES":
            resultFile = reportPath + "/(RedLine)" + str(policy[1]) + ".txt"
        else:
            resultFile = reportPath + "/" + str(policy[1]) + ".txt"
        if file[-5:]==".docx" or file[-4:]==".doc":
            result1 = scanWord(matchKey,file)
            if result1 != False:
                result = result + result1
        if file[-5:]==".xlsx" or file[-4:]==".xls":
            result2 = scanExcel(matchKey,file)
            if result2 != False:
                result = result + result2

        if not os.path.exists(reportPath):
            os.makedirs(reportPath)
        if result == []:
            continue
        f_result = open(resultFile,'a')
        for res in result:
            file = res[0]
            line = res[1]
            info = res[2]
            info = info.replace("\n","\\n")
            f_result.write(file+":"+line+":"+info+"\n")
        f_result.close()
    return 1

def scan_in_Document():
    global g_documentDir
    getScanPolicy()
    for conf in g_config:
        if conf[0]=="documentDir":
            g_documentDir = conf[1]
    if g_documentDir=="":
        g_Log.writeLog(u"错误信息：没有配置源码路径")
        return 0
    if g_scanPolicy==[]:
        g_Log.writeLog(u"错误信息：没有配置扫描策略")
        return 0
    reportPath = g_caseName.replace("TestCase","Report").replace(".py","")+"-"+g_Global.getValue("startTime").replace(":","").replace(" ","").replace("-","")
    if not os.path.exists(reportPath):
        os.makedirs(reportPath)
    shutil.copy(g_curDir+"/scanSensitivePolicy.xlsx",reportPath+"/scanSensitivePolicy.xlsx")

    for docFile in os.listdir(g_documentDir):
        tmpfile = os.path.join(g_documentDir,docFile).replace("\\","/")
        if tmpfile.split("/")[-1][0:2]=="~$":
            continue
        g_Log.writeLog("Begin to scan \"{file}\"".format(file=tmpfile))
        scanDocument(reportPath,tmpfile)



'''
""" 以下定义的函数，请在特定位置添加自己的代码 """
'''
# 执行前的准备操作
def prepare():
    try:
        ''''''''' 可以在此处下方添加自己的代码 '''''''''
    except:
        g_Log.writeLog("traceback")
        return 0
    return 1

# 执行用例
def run():
    try:
        ''''''''' 可以在此处下方添加自己的代码 '''''''''
        getScanPolicy()
        scan_in_Document()
    except:
        g_Log.writeLog("traceback")
        return 0
    return 1

# 执行后清理环境
def clearup():
    try:
        ''''''''' 可以在此处下方添加自己的代码 '''''''''
    except:
        g_Log.writeLog("traceback")
        return 0
    return 1

if __name__ == '__main__':
    res = prepare()
    if not res:
        print "执行用例prepare模块失败，结束用例{name}的执行".format(name=g_caseName)
    else:
        run()
        clearup()

