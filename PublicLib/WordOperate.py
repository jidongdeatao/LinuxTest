#-*- coding: utf-8 -*-

import os
import sys
import docx
import datetime
import shutil
from win32com.client import Dispatch,constants

reload(sys)
sys.setdefaultencoding('utf-8')

import GlobalValue as g_Global
import Log
g_Log = None
import LocalOperate
g_Local = None

class Word:
    def __init__(self,docName=None):
        global g_Log,g_Local
        g_Log = Log.Log()
        g_Local = LocalOperate.Local()
        '''
        docDir = os.path.split(docName)[0]
        docSuffix = docName.split(".")[-1]
        if g_Global.getValue("startTime") is None:
            g_Global.setValue("startTime",str(datetime.datetime.now()))
        tempDoc = docDir+"/"+"tempDoc-{time}.{suffix}".format(suffix=docSuffix,time=g_Global.getValue("startTime").replace(":",""))
        shutil.copy(docName,tempDoc)
        '''
        self.docName = docName

    def new(self):#新建一个word
        try:
            g_Log.writeLog("该功能暂未实现")
            return 1
            # 检查excel所在路径是否存在，如不存在则新建完整目录
            excelPath = os.path.split(self.excelName)[0]
            pathExist = os.path.exists(excelPath)
            if not pathExist:
                os.makedirs(excelPath)
        except:
            g_Log.writeLog("traceback")
            return 0

    def write(self,result,redLine=0):# 向word写入
        try:
            g_Log.writeLog("该功能暂未实现")
            return 1
            # 读取word
        except:
            g_Log.writeLog("traceback")
            return 0

    def readlines(self): #读取word内容
        output = []
        try:
            g_Log.writeLog(u"打开并尝试读取{doc}".format(doc=g_Local.unicode(self.docName)))
            docName = self.docName
            file = docx.Document(docName)
            for i in range(len(file.paragraphs)):
                output.append(file.paragraphs[i].text)
            g_Log.writeLog(u"读取成功")
            return output
        except:
            return 0









